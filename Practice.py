from docx import Document 
from docx.shared import Inches
import pyttsx3

def speak(text):
    pyttsx3.speak(text)

document = Document()

# profile picture
document.add_picture(
    'me1.png', 
    width=Inches(1.0)
    )


# name phone number and email details
name = input('What is your name? ')
speak('Hello ' + name + '. how is your today?')

speak('What is your phone number?')
phone_number = input('What is your phone number? ')

speak('What is your email?')
email = input('What is your E-mail? ')

document.add_paragraph(
    name + '  |  ' + phone_number + '  |  ' + email
)

# about me
document.add_heading('About me')

speak('Tell about yourself')
document.add_paragraph(
    input('Tell about yourself ')
)

# work experience
document.add_heading('Work Experience')

# more experiences
while True:
    has_more_experiences = input(
        'Do you have more experiences? Yes or No ' + ' '
    )
    if has_more_experiences.lower() == 'yes':
        p = document.add_paragraph()

        company = input('Enter company ')
        from_date = input('From Date ')
        to_date = input('To Date ')

        p.add_run(company + '  ').bold = True
        p.add_run(from_date + ' - ' + to_date + '\n').italic = True

        experience_details = input(
            'Decribe your experience at ' + company + ' '
        )
        p.add_run(experience_details)
        
    else:
        break
    
# skills
document.add_heading('Skills')

while True:
    has_more_skills = input(
        'Do you have more skills? Yes or No ' + ' '
    )
    if has_more_skills.lower() == 'yes':
        skill = input('Enter your skill ')
        p1 = document.add_paragraph(skill)
        
        p1.style = 'List Bullet'
        
    else:
        break
    
# footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = "CV generated using Python"
        

document.save('cv.docx')